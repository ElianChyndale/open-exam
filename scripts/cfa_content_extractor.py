"""C+ Content Extraction Pipeline

Extracts textbook body content from ePub, PDF, DOCX, and DOC source files,
maps it to KnowledgeBlock fields via rule-based heuristics.

Usage:
    python -c "from scripts.cfa_content_extractor import *; ..."

Architecture:
    SourceFormat (ePub/PDF/DOCX/DOC)
        → ExtractedSection (format-agnostic intermediate)
        → KnowledgeBlock fields (rule-based mapper)
        → C+ renderer integration
"""

from __future__ import annotations

import json
import logging
import re
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable
from xml.etree.ElementTree import Element

import lxml.etree as ET

logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")
log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

REPO_ROOT = Path(__file__).resolve().parent.parent
STRATEGY_ROOT = REPO_ROOT / ".system" / "memory" / "strategy"
CACHE_PATH = STRATEGY_ROOT / "cfa-content-extraction-cache.json"

# EPUB root — same as other scripts
EPUB_ROOT = Path(r"D:\BaiduNetdiskDownload\CFA2026一级原版书")

# ---------------------------------------------------------------------------
# Data Model
# ---------------------------------------------------------------------------


@dataclass
class ExtractedSection:
    """Format-agnostic container for one textbook section's extracted content.

    Every field is optional; the rule-based mapper works with whatever is
    available and falls back to sensible templates for missing fields.
    """

    section_number: str = ""
    heading: str = ""
    paragraphs: list[str] = field(default_factory=list)
    english_terms: list[str] = field(default_factory=list)
    formula_texts: list[str] = field(default_factory=list)
    definition_sentences: list[str] = field(default_factory=list)
    example_paragraphs: list[str] = field(default_factory=list)
    warning_sentences: list[str] = field(default_factory=list)
    procedural_sentences: list[str] = field(default_factory=list)
    los_text: str = ""
    source: str = ""


# ---------------------------------------------------------------------------
# ePub Extractor
# ---------------------------------------------------------------------------


class EpubSectionExtractor:
    """Extracts sections from CFA 2026 ePub files.

    Each volume ePub has:
      - OEBPS/nav.xhtml  → TOC with section → href mapping
      - OEBPS/*.xhtml    → actual body content (XHTML)
    """

    def __init__(self, epub_path: Path) -> None:
        if not epub_path.exists():
            raise FileNotFoundError(f"ePub not found: {epub_path}")
        self._epub_path = epub_path
        self._zip: zipfile.ZipFile | None = None
        # Cache: parsed nav entries [(title, href, section_id, file_in_zip), ...]
        self._nav_entries: list[tuple[str, str, str, str]] | None = None

    def __enter__(self) -> EpubSectionExtractor:
        self._zip = zipfile.ZipFile(self._epub_path)
        self._build_nav_index()
        return self

    def __exit__(self, *args: Any) -> None:
        if self._zip:
            self._zip.close()
            self._zip = None

    # ------------------------------------------------------------------
    # nav.xhtml parsing
    # ------------------------------------------------------------------

    def _build_nav_index(self) -> None:
        """Parse nav.xhtml → list of (title, href, section_id, file_in_zip)."""
        assert self._zip is not None
        entries: list[tuple[str, str, str, str]] = []
        try:
            nav_xml = ET.fromstring(
                self._zip.read("OEBPS/nav.xhtml"), parser=ET.HTMLParser()
            )
        except KeyError:
            log.warning("nav.xhtml not found in %s", self._epub_path.name)
            self._nav_entries = entries
            return

        # HTMLParser strips namespaces, so use plain tag names
        for a in nav_xml.iter("a"):
            href = a.get("href", "")
            title = "".join(a.itertext()).strip()
            if not title or not href or href.startswith("cover"):
                continue
            if ".xhtml" not in href:
                continue
            # Split href into file and anchor
            if "#" in href:
                file_part, anchor = href.split("#", 1)
            else:
                file_part, anchor = href, ""
            # Normalise file path — strip leading path elements
            file_in_zip = (
                file_part if file_part.startswith("OEBPS/") else f"OEBPS/{file_part}"
            )
            entries.append((title, href, anchor, file_in_zip))

        self._nav_entries = entries

    def match_section(self, signal_topic: str) -> tuple[str, str, str] | None:
        """Match a signal_topic string to a nav entry.

        Returns (file_in_zip, section_id, nav_title) or None.
        """
        if self._nav_entries is None:
            return None

        # Normalise the query
        query_num = self._extract_number(signal_topic)
        query_clean = self._clean_title(signal_topic)

        best: tuple[str, str, str] | None = None
        best_score = -1

        for title, href, anchor, fz in self._nav_entries:
            entry_num = self._extract_number(title)
            entry_clean = self._clean_title(title)

            # Exact number match (highest confidence)
            if query_num and entry_num and query_num == entry_num:
                score = 100
                # Bonus for exact title match
                if query_clean.lower() == entry_clean.lower():
                    score = 200
                if score > best_score:
                    best_score = score
                    best = (fz, anchor, title)

            # Clean title substring match
            if query_clean and entry_clean:
                if (
                    query_clean.lower() in entry_clean.lower()
                    or entry_clean.lower() in query_clean.lower()
                ):
                    score = 50 - abs(len(query_clean) - len(entry_clean))
                    if score > best_score:
                        best_score = score
                        best = (fz, anchor, title)

        return best

    # ------------------------------------------------------------------
    # Section content extraction
    # ------------------------------------------------------------------

    def extract_section(self, signal_topic: str) -> ExtractedSection:
        """Extract content for one signal section."""
        matched = self.match_section(signal_topic)
        if not matched:
            log.debug("No nav match for: %s", signal_topic)
            return ExtractedSection(
                heading=self._clean_title(signal_topic),
                paragraphs=[],
                english_terms=[],
            )

        fz, anchor, nav_title = matched
        assert self._zip is not None

        try:
            raw = self._zip.read(fz)  # bytes — lxml needs bytes for XML-declared content
        except KeyError:
            log.warning("File not found in epub: %s", fz)
            return ExtractedSection(heading=self._clean_title(signal_topic))

        return self._parse_xhtml_section(raw, anchor, signal_topic)

    def _parse_xhtml_section(
        self, xhtml: bytes, anchor: str, signal_topic: str
    ) -> ExtractedSection:
        """Parse one section's XHTML into an ExtractedSection."""
        root = ET.fromstring(xhtml, parser=ET.HTMLParser())
        sec = ExtractedSection(
            heading=self._clean_title(signal_topic),
            source=self._epub_path.name,
        )

        # Find the target element
        # Handle both id= and name= anchors
        target: Element | None = root.find(f".//*[@id='{anchor}']")
        if target is None:
            target = root.find(f".//*[@name='{anchor}']")
        if target is None:
            log.debug("Anchor %s not found in XHTML", anchor)
            sec.paragraphs = self._fallback_text_extract(root)
            return sec

        # Determine section boundaries
        section_number = self._extract_number(signal_topic)
        sec.section_number = section_number

        tag = target.tag.split("}")[-1] if "}" in target.tag else target.tag

        # If target is a container element (section, div), extract its children directly
        if tag in ("section", "div", "article"):
            for child in target:
                self._extract_from_element(child, sec)
            return sec

        # If target is a heading, collect siblings until next same-level heading
        parent = target.getparent() if target.getparent() is not None else root
        header_level = int(tag[1]) if tag.startswith("h") and len(tag) == 2 else 1

        iter_elements: list[Element] = []
        started = False
        for child in parent:
            child_tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if not started:
                if child is target or child.get("id", "") == anchor:
                    started = True
                else:
                    continue
            # Stop at next same-level heading (different section)
            if child is not target and child_tag.startswith("h"):
                child_lev = int(child_tag[1]) if len(child_tag) == 2 else 1
                if child_lev <= header_level:
                    break
            iter_elements.append(child)

        for elem in iter_elements:
            self._extract_from_element(elem, sec)

        return sec

    def _extract_terms(self, elem: Element, sec: ExtractedSection) -> None:
        """Extract dfn/em terms and MathML from ANY element (called for all tag types)."""
        for dfn in elem.iter("dfn"):
            term = "".join(dfn.itertext()).strip()
            if term and term not in sec.english_terms and len(term) > 1:
                sec.english_terms.append(term)
        for em in elem.iter("em"):
            term = "".join(em.itertext()).strip()
            if term and term not in sec.english_terms and len(term) > 1:
                sec.english_terms.append(term)
        for math in elem.iter("math"):
            text = self._render_math(math)
            if text and text not in sec.formula_texts:
                sec.formula_texts.append(text)

    def _extract_from_element(self, elem: Element, sec: ExtractedSection) -> None:
        """Extract text, terms, formulas from one XML element into sec."""
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

        # Always extract terms from this element
        self._extract_terms(elem, sec)

        # Skip LOS headings and sub-chapter markers
        if elem.get("class", "") in ("los", "overview-caption"):
            return
        if tag == "figure":
            self._extract_figure(elem, sec)
            return
        if tag in ("section", "article"):
            for child in elem:
                self._extract_from_element(child, sec)
            return

        if tag in ("h1", "h2", "h3", "h4"):
            return  # We already have the heading

        if tag == "p":
            text = "".join(elem.itertext()).strip()
            if not text or len(text) < 10:
                return
            text_clean = re.sub(r"\s+", " ", text).strip()

            # Detect formulas
            if self._is_formula_paragraph(elem):
                sec.formula_texts.append(text_clean)
                return
            # Detect examples
            if re.match(
                r"(example|suppose|consider|assume|for\s+instance)",
                text_clean,
                re.IGNORECASE,
            ):
                sec.example_paragraphs.append(text_clean)
                return
            # Detect warning/contrast
            if re.search(
                r"\b(however|but|note\s+that|important|caution|cannot|must\s+not|only\s+if|be\s+careful)\b",
                text_clean,
                re.IGNORECASE,
            ):
                sec.warning_sentences.append(text_clean)
                return
            # Detect procedural
            if re.search(
                r"\b(to\s+(calculate|determine|find|compute|evaluate)|first\s*[,:]|step\s+\d|procedure)\b",
                text_clean,
                re.IGNORECASE,
            ):
                sec.procedural_sentences.append(text_clean)
                return
            # Detect definition sentences
            if re.search(r"\bis\s+(the\s+)?(rate|return|measure|value|cost|price|amount|difference)", text_clean, re.IGNORECASE) or \
               re.search(r"\b(refers\s+to|is\s+defined\s+as|means\s+that)", text_clean, re.IGNORECASE):
                sec.definition_sentences.append(text_clean)
                return

            sec.paragraphs.append(text_clean)
            return

        if tag in ("ul", "ol", "li"):
            for child in elem:
                self._extract_from_element(child, sec)
            return

        if tag == "span":
            cls = elem.get("class", "")
            if "h1-label" in cls or "h2-label" in cls or "h3-label" in cls:
                return

    def _extract_figure(self, elem: Element, sec: ExtractedSection) -> None:
        """Extract from <figure> elements (usually overview boxes)."""
        figcaption = elem.find(".//figcaption")
        if figcaption is not None:
            text = "".join(figcaption.itertext()).strip()
            if text and "Overview" in text:
                # Overview box — extract all list items as definitions
                for li in elem.iter("li"):
                    p_text = "".join(li.itertext()).strip()

    def _render_math(self, math_elem: Element) -> str:
        """Render MathML to readable text."""
        parts: list[str] = []
        # Process mi, mn, mo, mrow elements
        for elem in math_elem.iter():
            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
            if tag in ("mi", "mn", "mo", "mtext", "msup", "msub", "mfrac", "msqrt"):
                text = "".join(elem.itertext()).strip()
                if tag == "mi" and text:
                    parts.append(text)
                elif tag == "mn" and text:
                    parts.append(text)
                elif tag == "mo":
                    # Map operators
                    op_map = {
                        "−": " - ",
                        "×": " * ",
                        "÷": " / ",
                        "±": " +/- ",
                        "≤": " <= ",
                        "≥": " >= ",
                        "≠": " != ",
                        "≈": " ~ ",
                        "∑": " SUM ",
                        "∏": " PRODUCT ",
                        "√": " sqrt ",
                    }
                    mo = text.strip()
                    parts.append(op_map.get(mo, mo))
        result = "".join(parts)
        result = re.sub(r"\s+", " ", result).strip()
        return result

    def _is_formula_paragraph(self, elem: Element) -> bool:
        """Check if a paragraph is primarily a formula."""
        # Has MathML inside (HTMLParser strips namespace → just look for "math")
        if elem.find(".//math") is not None or "math" in str(elem.tag).lower():
            return True
        text = "".join(elem.itertext()).strip()
        # Contains formula patterns
        if re.search(r"^[A-Za-z]\s*[=+−]", text):
            return True
        if re.search(r"[=+−÷×]\s*[A-Za-z0-9]", text) and len(text) < 120:
            return True
        return False

    def _fallback_text_extract(self, root: Element) -> list[str]:
        """Fallback: extract all text paragraphs when anchor not found."""
        paragraphs: list[str] = []
        for p in root.iter("p"):
            text = "".join(p.itertext()).strip()
            if text and len(text) > 20:
                paragraphs.append(re.sub(r"\s+", " ", text).strip())
        return paragraphs

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_number(title: str) -> str:
        m = re.match(r"^(\d+(?:\.\d+)*)", title.strip())
        return m.group(1) if m else ""

    @staticmethod
    def _clean_title(title: str) -> str:
        return re.sub(r"^\d+(?:\.\d+)*\.?\s*", "", title).strip()


# ---------------------------------------------------------------------------
# PDF Extractor  (fallback when ePub not available)
# ---------------------------------------------------------------------------


class PdfSectionExtractor:
    """Extract sections from CFA PDF files using pdfplumber.

    PDF extraction is lower quality than ePub; this is a fallback.
    Section detection works by scanning pages for heading-like lines.
    """

    def __init__(self, pdf_path: Path) -> None:
        if not pdf_path.exists():
            raise FileNotFoundError(f"PDF not found: {pdf_path}")
        self._pdf_path = pdf_path
        self._heading_pattern = re.compile(
            r"^\d+(?:\.\d+)*\.?\s+[A-Z][A-Za-z\s,/:]+$"
        )

    def extract_section(self, signal_topic: str) -> ExtractedSection:
        """Extract content for a signal topic from PDF."""
        import pdfplumber

        sec = ExtractedSection(
            heading=self._clean_title(signal_topic),
            source=self._pdf_path.name,
        )
        query_clean = self._clean_title(signal_topic).lower()
        query_num = self._extract_number(signal_topic)

        with pdfplumber.open(self._pdf_path) as pdf:
            found = False
            collected: list[str] = []
            for page in pdf.pages:
                text = page.extract_text() or ""
                lines = text.splitlines()
                for line in lines:
                    line = line.strip()
                    if not line or len(line) < 5:
                        continue
                    # Check if this line is a heading
                    line_clean = self._clean_title(line).lower()
                    line_num = self._extract_number(line)

                    if found:
                        # Stop at next same-level heading
                        if self._is_heading(line) and line_num:
                            next_parent = query_num.split(".")[0] if query_num else ""
                            this_parent = line_num.split(".")[0] if line_num else ""
                            if this_parent != next_parent:
                                break
                            # Sub-headings of our section are OK
                            if line_num.count(".") <= query_num.count(".") and line_num != query_num:
                                break
                        collected.append(line)
                    elif query_clean in line_clean or (query_num and query_num == line_num):
                        found = True
                        collected.append(line)

                if found and not text:
                    break  # Heading match found but no more content

            # Classify collected lines
            if collected:
                sec.paragraphs = [
                    re.sub(r"\s+", " ", l).strip()
                    for l in collected
                    if len(l) > 30
                ]
                # Extract terms (italic/bold markers are lost in PDF text extraction)
                for p in sec.paragraphs:
                    # Try to extract italic-looking terms (surrounded by special chars)
                    terms = re.findall(r"_([A-Za-z][A-Za-z\s,-]+)_", p)
                    for t in terms:
                        t = t.strip()
                        if t and t not in sec.english_terms:
                            sec.english_terms.append(t)

        return sec

    @staticmethod
    def _is_heading(line: str) -> bool:
        # Heuristic: short lines starting with a number followed by capital letter
        return bool(re.match(r"^\d+(?:\.\d+)*\.[A-Z]", line)) or bool(
            re.match(r"^\d+(?:\.\d+)*\s+[A-Z]", line)
        )

    @staticmethod
    def _extract_number(title: str) -> str:
        m = re.match(r"^(\d+(?:\.\d+)*)", title.strip())
        return m.group(1) if m else ""

    @staticmethod
    def _clean_title(title: str) -> str:
        return re.sub(r"^\d+(?:\.\d+)*\.?\s*", "", title).strip()


# ---------------------------------------------------------------------------
# DOCX Extractor
# ---------------------------------------------------------------------------


class DocxSectionExtractor:
    """Extract sections from DOCX files using python-docx.

    DOCX has heading styles (Heading 1, Heading 2, ...) that make
    section boundary detection reliable.
    """

    def __init__(self, docx_path: Path) -> None:
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX not found: {docx_path}")
        self._path = docx_path

    def extract_section(self, signal_topic: str) -> ExtractedSection:
        from docx import Document

        sec = ExtractedSection(
            heading=self._clean_title(signal_topic),
            source=self._path.name,
        )
        query_clean = self._clean_title(signal_topic).lower()
        query_num = self._extract_number(signal_topic)

        doc = Document(str(self._path))
        found = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else ""

            if found:
                # Stop at next heading of same or higher level
                if style.startswith("Heading") and text.strip():
                    line_num = self._extract_number(text)
                    if line_num and query_num:
                        if line_num.count(".") <= query_num.count("."):
                            break
                # Classify
                runs_bold = any(r.bold for r in para.runs if r.bold)
                runs_italic = any(r.italic for r in para.runs if r.italic)
                if runs_bold and len(text) > 3:
                    sec.english_terms.append(text)
                if runs_italic and len(text) > 3:
                    if text not in sec.english_terms:
                        sec.english_terms.append(text)
                if style.startswith("Heading") or len(text) > 30:
                    sec.paragraphs.append(text)

            elif query_clean in text.lower() or (
                query_num and self._extract_number(text) == query_num
            ):
                found = True
                sec.paragraphs.append(text)

        return sec

    @staticmethod
    def _extract_number(title: str) -> str:
        m = re.match(r"^(\d+(?:\.\d+)*)", title.strip())
        return m.group(1) if m else ""

    @staticmethod
    def _clean_title(title: str) -> str:
        return re.sub(r"^\d+(?:\.\d+)*\.?\s*", "", title).strip()


# ---------------------------------------------------------------------------
# DOC Extractor  (converts DOC → DOCX first via LibreOffice)
# ---------------------------------------------------------------------------


class DocSectionExtractor:
    """Converts .doc to .docx using LibreOffice, then delegates to DocxSectionExtractor.

    Requires LibreOffice installed at a known path.
    """

    LIBREOFFICE_PATHS = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/lib/libreoffice/program/soffice",
        "/usr/bin/soffice",
    ]

    def __init__(self, doc_path: Path) -> None:
        if not doc_path.exists():
            raise FileNotFoundError(f"DOC not found: {doc_path}")
        self._doc_path = doc_path

    def _locate_soffice(self) -> str | None:
        for path in self.LIBREOFFICE_PATHS:
            if Path(path).exists():
                return path
        # Try PATH
        try:
            result = subprocess.run(
                ["soffice", "--version"],
                capture_output=True,
                text=True,
                timeout=10,
            )
            if result.returncode == 0:
                return "soffice"
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        return None

    def extract_section(self, signal_topic: str) -> ExtractedSection:
        soffice = self._locate_soffice()
        if soffice is None:
            log.warning(
                "LibreOffice not found; cannot convert %s. Install LibreOffice or "
                "convert manually.",
                self._doc_path.name,
            )
            return ExtractedSection(heading=self._clean_title(signal_topic))

        with tempfile.TemporaryDirectory() as tmp:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "docx",
                 "--outdir", tmp, str(self._doc_path)],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode != 0:
                log.warning("LibreOffice conversion failed: %s", result.stderr)
                return ExtractedSection(heading=self._clean_title(signal_topic))

            tmp_path = Path(tmp)
            docx_files = list(tmp_path.glob("*.docx"))
            if not docx_files:
                return ExtractedSection(heading=self._clean_title(signal_topic))

            # Delegate to DOCX extractor
            docx_ext = DocxSectionExtractor(docx_files[0])
            return docx_ext.extract_section(signal_topic)

    @staticmethod
    def _clean_title(title: str) -> str:
        return re.sub(r"^\d+(?:\.\d+)*\.?\s*", "", title).strip()

    @staticmethod
    def _extract_number(title: str) -> str:
        m = re.match(r"^(\d+(?:\.\d+)*)", title.strip())
        return m.group(1) if m else ""


# ---------------------------------------------------------------------------
# Extractor Registry
# ---------------------------------------------------------------------------


def _guess_source_format(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".epub":
        return "epub"
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext == ".doc":
        return "doc"
    return "unknown"


def get_extractor(
    source_path: Path,
) -> EpubSectionExtractor | PdfSectionExtractor | DocxSectionExtractor | DocSectionExtractor | None:
    """Return the appropriate extractor for a source file."""
    fmt = _guess_source_format(source_path)
    if fmt == "epub":
        return EpubSectionExtractor(source_path)
    elif fmt == "pdf":
        return PdfSectionExtractor(source_path)
    elif fmt == "docx":
        return DocxSectionExtractor(source_path)
    elif fmt == "doc":
        return DocSectionExtractor(source_path)
    else:
        log.warning("Unsupported source format: %s", source_path.suffix)
        return None


# ---------------------------------------------------------------------------
# Rule-Based KnowledgeBlock Field Mapper
# ---------------------------------------------------------------------------


def extract_core_meaning(sec: ExtractedSection) -> str:
    """Extract core meaning: definition sentences first, then first paragraphs."""
    if sec.definition_sentences:
        return sec.definition_sentences[0][:500]

    # Look for a paragraph that defines/introduces the concept
    for p in sec.paragraphs:
        if any(
            marker in p.lower()[:100]
            for marker in [" is ", " are ", " refers ", " defined ", " means "]
        ):
            return p[:500]

    # Fallback: first substantive paragraph
    for p in sec.paragraphs:
        if len(p) > 50:
            return p[:500]

    return f"本节讲解 {sec.heading} 的定义和应用。" if sec.heading else ""


def _is_real_term(term: str) -> bool:
    """Check if a string looks like a real technical term, not a sentence fragment."""
    t = term.strip().rstrip(".,;: ")
    if len(t) < 3:
        return False
    # Skip if it looks like a full sentence (capitalized first word + verb structure)
    if len(t) > 50 and " " in t and t[0].isupper():
        # Likely a sentence, not a term
        return False
    # Skip common false positives (regardless of case)
    skip_words = {"using", "suppose", "example", "equation", "figure", "table",
                  "illustrate", "demonstrate", "describe", "note that", "step",
                  "first", "second", "third", "therefore", "however", "because",
                  "the sum", "the product", "this is", "in this", "the following",
                  "consider", "assume", "given", "where:", "thus", "we", "the",
                  "an", "a"}
    if any(t.lower().startswith(w) for w in skip_words):
        return False
    # Accept multi-word terms even if lowercase (e.g. "opportunity cost")
    # Accept single capitalized or all-caps terms
    if t[0].islower() and " " not in t and len(t) < 6:
        return False  # single lowercased short word like "the", "rate"
    return True


def extract_english_terms(sec: ExtractedSection) -> str:
    """Combine dfn, em, and likely technical terms."""
    terms: list[str] = []
    seen: set[str] = set()

    for t in sec.english_terms:
        key = t.lower().strip().rstrip(".,;: ")
        if key not in seen and _is_real_term(t):
            seen.add(key)
            terms.append(t.strip().rstrip(".,;: "))

    # Scan definition sentences for capitalized multi-word terms
    for s in sec.definition_sentences:
        for m in re.findall(r"\b[A-Z][a-z]+(?:\s+[a-z]+)*\s+[A-Z][a-zA-Z]+", s):
            key = m.lower()
            if key not in seen and _is_real_term(m):
                seen.add(key)
                terms.append(m.strip())

    return "; ".join(terms[:12]) if terms else sec.heading


def extract_formula_rule(sec: ExtractedSection) -> str:
    """Build formula rule from MathML and formula paragraphs."""
    formulas: list[str] = list(sec.formula_texts)

    # Also check paragraphs for formula lines
    for p in sec.paragraphs:
        if any(op in p for op in ["=", "+", "−", "×", "÷", "∑", "∏", "√"]):
            if len(p) < 150:
                formulas.append(p.strip())

    if formulas:
        return "；".join(formulas[:3])

    # Fallback: look for formula patterns in definition sentences
    for s in sec.definition_sentences:
        if "=" in s:
            return s[:200]

    return ""


def extract_why_it_matters(sec: ExtractedSection) -> str:
    """Extract "why it matters" from paragraphs."""
    candidates: list[str] = []

    for p in sec.paragraphs:
        lower = p.lower()
        if any(
            marker in lower
            for marker in [
                "important", "key ", "critical", "essential",
                "because", "useful for", "used to", "allows",
                "enables", "primary", "fundamental",
            ]
        ):
            candidates.append(p)

    # Also check definition sentences
    for s in sec.definition_sentences:
        lower = s.lower()
        if any(marker in lower for marker in ["useful", "important", "key "]):
            candidates.append(s)

    if candidates:
        return candidates[0][:400]

    # Fallback: connect to exam
    return (
        f"理解 {sec.heading} 能帮助你准确判断考试中涉及此概念的题目类型和适用条件。"
        if sec.heading
        else ""
    )


def extract_exam_translation(sec: ExtractedSection) -> str:
    """Translate textbook content → exam action."""
    # Procedural sentences are the best source
    if sec.procedural_sentences:
        return sec.procedural_sentences[0][:400]

    # Definition sentences with action verbs
    for s in sec.definition_sentences:
        if any(v in s.lower() for v in ["calculate", "determine", "measure", "compute"]):
            return s[:400]

    # Fallback: construct from heading
    if sec.heading:
        # Try to create exam translation from the term names
        terms = extract_english_terms(sec)[:80]
        return (
            f"考试中会围绕「{sec.heading}」出题。做题时先判断题目要你识别、计算还是解释，"
            f"然后找到对应的概念边界和公式。识别关键术语：{terms}"
        )
    return ""


def extract_question_triggers(sec: ExtractedSection) -> str:
    """Extract trigger words that appear in exam questions."""
    triggers: list[str] = []

    # English terms are the primary triggers
    terms = extract_english_terms(sec)
    if terms:
        triggers.append(f"题干出现 `{terms[:60]}`")

    # Add heading word triggers
    if sec.heading:
        words = re.findall(r"\b[A-Z][a-z]+\b", sec.heading)
        for w in words[:3]:
            triggers.append(f"`{w}`")

    # Add formula-related triggers
    if sec.formula_texts:
        # Extract variable names
        vars_found = re.findall(r"\b[A-Z][A-Za-z0-9]*\b", " ".join(sec.formula_texts[:2]))
        for v in vars_found[:3]:
            v = v.strip()
            if v and v not in ("The", "This", "That", "For", "SUM", "PRODUCT", "sqrt"):
                triggers.append(f"`{v}`")

    if triggers:
        return "；".join(triggers[:6])

    if sec.heading:
        return f"`{sec.heading}` 相关术语和概念"
    return ""


def extract_trap_fix_rule(sec: ExtractedSection) -> str:
    """Extract trap/fix rules from warning/contrast sentences."""
    if sec.warning_sentences:
        # Combine first warning with its context
        result = sec.warning_sentences[0][:400]
        if len(sec.warning_sentences) > 1:
            result += " | " + sec.warning_sentences[1][:200]
        return result

    # Look for "however", "but" in paragraphs
    for p in sec.paragraphs:
        if "however" in p.lower() or re.search(r"\bbut\b", p.lower()):
            return p[:400]
        if "note that" in p.lower():
            return p[:400]

    # Fallback
    terms = extract_english_terms(sec)[:60]
    if terms:
        return f"注意区分 `{terms[:60]}` 与其他相似概念的边界。考试陷阱常设置在同一上下文里混淆它们。"
    return ""


def build_knowledge_block_from_extraction(
    sec: ExtractedSection,
    textbook_position: str,
    formula_rule_fallback: str = "",
    links: str = "",
) -> dict[str, str]:
    """Build the 9 KnowledgeBlock fields from an ExtractedSection.

    Returns dict in the same shape as KnowledgeBlock dataclass.
    """
    return {
        "core_meaning": extract_core_meaning(sec) or "见教材原文。",
        "english_terms": extract_english_terms(sec) or sec.heading,
        "why_it_matters": extract_why_it_matters(sec) or "见教材原文。",
        "formula_rule": extract_formula_rule(sec) or formula_rule_fallback or "见教材原文公式。",
        "exam_translation": extract_exam_translation(sec) or "见教材原文，按 LOS 动词准备。",
        "question_triggers": extract_question_triggers(sec) or "见教材原文关键词。",
        "practice_mock_evidence": "基础题用于检验本概念的定义和应用场景。",
        "trap_fix_rule": extract_trap_fix_rule(sec) or "注意区分相似概念。",
    }


# ---------------------------------------------------------------------------
# Module-Level Extraction
# ---------------------------------------------------------------------------


def extract_module_sections(
    subject: str,
    module_module: str,
    signal_topics: list[str],
) -> dict[str, ExtractedSection]:
    """Extract all sections for a given module from the best available source.

    Priority: ePub → PDF → DOCX → DOC
    """
    from scripts.cfa_c_plus_redesign import load_textbook_index

    index = load_textbook_index()
    subject_index = next((s for s in index if s["subject"] == subject), None)
    if subject_index is None:
        log.warning("Subject %s not found in index", subject)
        return {}

    epub_path_str: str = subject_index.get("epub", "")
    if not epub_path_str:
        log.warning("No epub path for %s", subject)
        return _try_pdf_extraction(subject, signal_topics)

    epub_path = Path(epub_path_str)
    if not epub_path.exists():
        log.warning("ePub not found at %s; trying PDF fallback", epub_path)
        return _try_pdf_extraction(subject, signal_topics)

    result: dict[str, ExtractedSection] = {}
    try:
        with EpubSectionExtractor(epub_path) as ext:
            for signal in signal_topics:
                sec = ext.extract_section(signal)
                result[signal] = sec
                if sec.paragraphs or sec.definition_sentences:
                    log.debug(
                        "  ✓ %s (%d paras, %d terms, %d formulas)",
                        signal,
                        len(sec.paragraphs) + len(sec.definition_sentences),
                        len(sec.english_terms),
                        len(sec.formula_texts),
                    )
                else:
                    log.debug("  - %s (no content extracted)", signal)
    except Exception as exc:
        log.warning("ePub extract failed for %s %s: %s", subject, module_module, exc)
        return _try_pdf_extraction(subject, signal_topics)

    return result


def _try_pdf_extraction(
    subject: str, signal_topics: list[str]
) -> dict[str, ExtractedSection]:
    """Fallback: extract from PDF matching the subject volume."""
    pdf_path = _find_pdf_for_subject(subject)
    if pdf_path is None:
        return {}
    try:
        ext = PdfSectionExtractor(pdf_path)
        return {signal: ext.extract_section(signal) for signal in signal_topics}
    except Exception as exc:
        log.warning("PDF extract failed for %s: %s", subject, exc)
        return {}


def _find_pdf_for_subject(subject: str) -> Path | None:
    """Find a PDF in EPUB_ROOT that matches the subject."""
    subject_short = subject.split("(")[0].strip().replace(" ", "")
    pdfs = sorted(EPUB_ROOT.glob(f"*{subject_short}*.pdf"))
    if not pdfs:
        # Fallback: try volume-based matching
        from scripts.cfa_c_plus_redesign import load_textbook_index

        index = load_textbook_index()
        subject_index = next((s for s in index if s["subject"] == subject), None)
        if subject_index:
            epub_name = Path(subject_index.get("epub", "")).stem
            pdfs = sorted(EPUB_ROOT.glob(f"{epub_name}.pdf"))
    return pdfs[0] if pdfs else None


# ---------------------------------------------------------------------------
# Cache Management
# ---------------------------------------------------------------------------


def load_cache() -> dict[str, Any]:
    if CACHE_PATH.exists():
        try:
            return json.loads(CACHE_PATH.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            pass
    return {}


def save_cache(cache: dict[str, Any]) -> None:
    CACHE_PATH.parent.mkdir(parents=True, exist_ok=True)
    CACHE_PATH.write_text(
        json.dumps(cache, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    log.info("Cache saved: %s (%d subjects)", CACHE_PATH.name, len(cache))


def cache_key(subject: str, module: str) -> str:
    return f"{subject}::{module}"


def extract_and_cache(subject: str, module_module: str, signal_topics: list[str]) -> dict[str, ExtractedSection]:
    """Extract sections, cache them, return results."""
    cache = load_cache()
    key = cache_key(subject, module_module)

    if key in cache:
        log.info("Cache hit: %s", key)
        # Deserialize from cache
        return _deserialize_cache_entries(cache[key])

    log.info("Extracting: %s %s (%d sections)", subject, module_module, len(signal_topics))
    sections = extract_module_sections(subject, module_module, signal_topics)

    # Serialize to cache
    cache[key] = _serialize_cache_entries(sections)
    save_cache(cache)

    return sections


def _serialize_cache_entries(sections: dict[str, ExtractedSection]) -> dict[str, dict[str, Any]]:
    return {
        signal: {
            "section_number": sec.section_number,
            "heading": sec.heading,
            "paragraphs": sec.paragraphs,
            "english_terms": sec.english_terms,
            "formula_texts": sec.formula_texts,
            "definition_sentences": sec.definition_sentences,
            "example_paragraphs": sec.example_paragraphs,
            "warning_sentences": sec.warning_sentences,
            "procedural_sentences": sec.procedural_sentences,
            "los_text": sec.los_text,
            "source": sec.source,
        }
        for signal, sec in sections.items()
    }


def _deserialize_cache_entries(data: dict[str, dict[str, Any]]) -> dict[str, ExtractedSection]:
    return {
        signal: ExtractedSection(**fields)
        for signal, fields in data.items()
    }


def extract_all_subjects() -> int:
    """Extract all subjects and modules, cache results. Returns section count."""
    from scripts.cfa_c_plus_redesign import load_textbook_index

    index = load_textbook_index()
    total = 0
    for subject_index in index:
        subject = subject_index["subject"]
        for module in subject_index.get("modules", []):
            signals = module.get("signal_topics", [])
            if not signals:
                continue
            extract_and_cache(subject, module["module"], signals)
            total += len(signals)
    log.info("Extraction complete: %d sections cached", total)
    return total


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main() -> None:
    import argparse

    parser = argparse.ArgumentParser(description="C+ Content Extraction Pipeline")
    parser.add_argument("--extract-all", action="store_true", help="Extract all subjects and cache")
    parser.add_argument(
        "--subject", type=str, default="",
        help="Single subject to extract (e.g. 'Quantitative Methods')",
    )
    parser.add_argument("--module", type=str, default="", help="Single module (e.g. 'M01')")
    args = parser.parse_args()

    if args.extract_all:
        total = extract_all_subjects()
        print(f"Extracted {total} sections total")
        return

    if args.subject and args.module:
        from scripts.cfa_c_plus_redesign import load_textbook_index

        index = load_textbook_index()
        subject_index = next((s for s in index if s["subject"] == args.subject), None)
        if subject_index is None:
            print(f"Subject not found: {args.subject}")
            return
        module_idx = next(
            (m for m in subject_index.get("modules", []) if m["module"] == args.module),
            None,
        )
        if module_idx is None:
            print(f"Module not found: {args.module}")
            return

        signals = module_idx.get("signal_topics", [])
        sections = extract_and_cache(args.subject, args.module, signals)
        print(f"Extracted {len(sections)} sections for {args.subject} {args.module}")
        for signal, sec in sections.items():
            print(f"  {signal}: {len(sec.paragraphs)}p + {len(sec.definition_sentences)}d + "
                  f"{len(sec.english_terms)}t + {len(sec.formula_texts)}f")
        return

    parser.print_help()


if __name__ == "__main__":
    main()
